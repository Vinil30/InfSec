from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "Infant_Monitoring_Feature_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    paragraph.paragraph_format.space_after = Pt(0)


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    set_paragraph_shading(paragraph, "EEF3F8")


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_feature(doc, number, title, where, steps, identify, files):
    heading = doc.add_heading(f"{number}. {title}", level=2)
    heading.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    widths = (Inches(1.35), Inches(5.75))
    labels = ["Where", "Steps", "What to identify", "Relevant files"]
    values = [where, steps, identify, files]

    for row_idx, label in enumerate(labels):
        row = table.rows[row_idx]
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        set_cell_shading(row.cells[0], "DCEAF8")
        set_cell_text(row.cells[0], label, bold=True, color=(22, 61, 105))
        set_cell_text(row.cells[1], values[row_idx])

    doc.add_paragraph()


def add_feature_map(doc):
    doc.add_heading("Quick Feature Map", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Feature", "Primary file", "Purpose"]
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "143D69")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=(255, 255, 255))

    rows = [
        ("Backend entry point", "app.py", "Flask routes, frame decoding, full CV pipeline orchestration."),
        ("Frontend UI", "templates/index.html", "Page structure, status panel, toolbar, and overlay container."),
        ("Frontend logic", "static/script.js", "Webcam capture, frame uploads, drawing overlays, alarm handling."),
        ("Styling", "static/style.css", "Responsive monitoring layout and risk visual states."),
        ("YOLO person detection", "utils/detector.py", "Pretrained YOLOv8 nano person detection."),
        ("ByteTrack infant tracking", "utils/tracker.py", "Selected infant track lock and tracking status."),
        ("Geofence", "utils/geofence.py", "Polygon normalization and centroid-in-region checks."),
        ("MediaPipe Pose", "utils/pose_estimator.py", "Pose keypoints after fence breach."),
        ("Object analysis", "utils/object_detector.py", "Nearby object detection and Safe/Warning/Dangerous categories."),
        ("Risk rules", "utils/risk_analyzer.py", "LOW/MEDIUM/HIGH scoring from fence, pose, and object signals."),
        ("Alerts", "utils/alert.py", "Alarm and warning payloads for the frontend."),
    ]
    for feature, file_name, purpose in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], feature)
        set_cell_text(cells[1], file_name)
        set_cell_text(cells[2], purpose)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.color.rgb = RGBColor(20, 61, 105)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 86, 135)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Powered Infant Monitoring Web App")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(20, 61, 105)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Step-by-step feature identification guide")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(92, 105, 120)

    intro = doc.add_paragraph()
    intro.add_run("Project location: ").bold = True
    intro.add_run(r"C:\Users\VINIL\Documents\AURA")

    doc.add_heading("Run The Application", level=1)
    doc.add_paragraph("Install dependencies if needed:")
    add_code_block(doc, "pip install -r requirements.txt")
    doc.add_paragraph("Start the Flask server from the project folder:")
    add_code_block(doc, "python app.py")
    doc.add_paragraph("Open the local web app:")
    add_code_block(doc, "http://127.0.0.1:5000")

    doc.add_heading("Feature Walkthrough", level=1)
    features = [
        (
            "Start Webcam",
            "Browser page, bottom toolbar.",
            "Click Start Camera and allow webcam access.",
            "The live webcam feed appears in the main video area. This confirms browser webcam capture and frame acquisition.",
            "templates/index.html, static/script.js",
        ),
        (
            "YOLO Person Detection",
            "Main video overlay after the camera starts.",
            "No extra action is needed. Detection begins automatically.",
            "Blue boxes labeled Person appear around detected people. These come from pretrained YOLOv8 nano.",
            "utils/detector.py, app.py",
        ),
        (
            "Select The Infant",
            "Bottom toolbar and main video area.",
            "Click Select Infant, then click once on the infant in the live feed.",
            "The selected infant receives a green box. The Tracking status changes from Waiting to a track ID.",
            "static/script.js, utils/tracker.py",
        ),
        (
            "ByteTrack Tracking",
            "Right status panel and green infant overlay box.",
            "Select the infant once. The backend keeps the selected track active afterward.",
            "Tracking displays a track number and the green box follows the selected infant.",
            "utils/tracker.py",
        ),
        (
            "Geofence Safe Region",
            "Main video overlay and bottom toolbar.",
            "Use the default yellow safe region, or click Draw Fence, place at least three points, then click Close Fence.",
            "A yellow polygon marks the safe area. The Fence status shows Inside while the infant centroid remains inside.",
            "utils/geofence.py, static/script.js",
        ),
        (
            "Fence Breach Detection",
            "Right status panel, alerts section, and video overlay.",
            "Move the selected infant outside the yellow geofence.",
            "Fence changes to Breached. Risk can rise from LOW to MEDIUM, and pose/object analysis becomes active.",
            "utils/geofence.py, utils/risk_analyzer.py",
        ),
        (
            "Pose Detection Phase",
            "Main video overlay after geofence breach.",
            "Trigger a fence breach. Pose detection runs only after that point.",
            "White pose keypoints may appear. Risk rules inspect head position, torso angle, and fall-like posture.",
            "utils/pose_estimator.py, utils/risk_analyzer.py",
        ),
        (
            "Surroundings Object Detection",
            "Main video overlay and Surroundings panel.",
            "Trigger a fence breach and keep nearby objects visible in the webcam frame.",
            "Nearby objects get labeled boxes and are categorized as Safe, Warning, or Dangerous.",
            "utils/object_detector.py, static/script.js",
        ),
        (
            "Object Distance And Overlap Risk",
            "Alerts section and object overlays after a fence breach.",
            "Place a COCO-detectable dangerous object near or overlapping the infant box.",
            "Close dangerous objects increase risk sharply. Overlapping dangerous objects trigger HIGH risk immediately.",
            "utils/object_detector.py, utils/risk_analyzer.py",
        ),
        (
            "Risk Level Display",
            "Right panel, Risk field and Alerts section.",
            "Watch the risk value as the infant stays inside, breaches the fence, or approaches risky objects.",
            "LOW is green, MEDIUM is yellow, HIGH is red. Alert reasons explain the rule decisions.",
            "utils/risk_analyzer.py",
        ),
        (
            "High Risk Alarm",
            "Full video area and browser audio.",
            "Create a HIGH risk condition through dangerous overlap or combined risk rules.",
            "A red HIGH RISK overlay appears and the browser alarm starts.",
            "utils/alert.py, static/script.js",
        ),
        (
            "Reset Everything",
            "Bottom toolbar.",
            "Click Reset.",
            "Infant selection, geofence, alarm, tracker state, and backend risk state are cleared.",
            "app.py, static/script.js",
        ),
    ]

    for idx, feature in enumerate(features, start=1):
        add_feature(doc, idx, *feature)

    add_feature_map(doc)

    doc.add_heading("Important MVP Notes", level=1)
    notes = [
        "No custom models are trained. The application uses pretrained YOLOv8 nano and MediaPipe Pose.",
        "Surroundings analysis intentionally starts only after the geofence is breached to reduce unnecessary computation.",
        "COCO class names are used for the MVP, so some real hazards are approximated through available classes such as bottle, knife, scissors, and electronics.",
        "This is a safety-support MVP and not a certified medical or child-safety system. Direct supervision is still required.",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
